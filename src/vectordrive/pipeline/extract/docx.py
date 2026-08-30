"""DOCX extraction via python-docx.

DOCX has no fixed "pages" (pagination is a rendering-time concern in Word,
not stored in the file). For v0.1 we treat the whole document as a single
logical page (page_number=1) — simple, and citations will still say
"page 1" honestly rather than fabricate page boundaries python-docx cannot
give us.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from vectordrive.pipeline.extract.base import ExtractedDocument, ExtractedPage


class DocxExtractor:
    NAME = "python-docx"
    VERSION = "1"

    def extract(self, path: Path) -> ExtractedDocument:
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        page = ExtractedPage(page_number=1, text=text, text_source="native", needs_ocr=False)
        return ExtractedDocument(
            extractor_name=self.NAME, extractor_version=self.VERSION, page_count=1, pages=[page]
        )
