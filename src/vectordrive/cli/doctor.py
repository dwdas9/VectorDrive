"""`vectordrive doctor`: environment and dependency smoke tests.

Each check is a small, independently-testable function returning a
CheckResult. `run_doctor()` runs them all and persists the vector backend
that `check_sqlite_vec_smoke_test()` determines into config.toml, so
`index`/`search` (M8+) read a stable, already-decided choice instead of
re-probing on every run.
"""
from __future__ import annotations

import io
import sqlite3
import sys
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

import httpx

from vectordrive.config import Config, load_config, path_is_inside, save_config


class CheckStatus(str, Enum):
    PASS = "PASS"
    WARN = "WARN"
    FAIL = "FAIL"


@dataclass
class CheckResult:
    name: str
    status: CheckStatus
    message: str
    required: bool = True


def check_python_and_sqlite() -> CheckResult:
    """Python 3.12.x whose sqlite3 module supports loadable extensions.

    python.org installers and Apple's system python3 build sqlite3 without
    enable_load_extension(); only Homebrew (and Conda) builds tested so far
    support it. sqlite-vec cannot load without this.
    """
    version = f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}"
    conn = sqlite3.connect(":memory:")
    try:
        conn.enable_load_extension(True)
        conn.enable_load_extension(False)
    except AttributeError:
        return CheckResult(
            "python_sqlite_extensions",
            CheckStatus.FAIL,
            f"Python {version}'s sqlite3 module does not support "
            "enable_load_extension(), which sqlite-vec requires. Fix: "
            "`brew install python@3.12` and recreate .venv with "
            "/opt/homebrew/bin/python3.12 — python.org installers and "
            "Apple's system Python lack this.",
        )
    finally:
        conn.close()

    if sys.version_info[:2] != (3, 12):
        return CheckResult(
            "python_sqlite_extensions",
            CheckStatus.WARN,
            f"Running Python {version} (expected 3.12.x). sqlite3 extension "
            "loading works, so continuing, but this isn't the pinned interpreter.",
        )
    return CheckResult(
        "python_sqlite_extensions",
        CheckStatus.PASS,
        f"Python {version}; sqlite3 supports loadable extensions.",
    )


def check_sqlite_vec_smoke_test() -> tuple[CheckResult, str]:
    """Real Apple Silicon smoke test: load sqlite-vec, create a vec0 table,
    insert a vector and run a KNN query. Falls back to the (as-yet
    unimplemented) NumPy backend on any failure, per spec.

    Returns (CheckResult, vector_backend).
    """
    try:
        import sqlite_vec
    except ImportError as exc:
        return (
            CheckResult(
                "sqlite_vec_smoke_test",
                CheckStatus.WARN,
                f"sqlite-vec not importable ({exc}); falling back to the "
                "NumPy vector-search backend.",
                required=False,
            ),
            "numpy",
        )

    conn = sqlite3.connect(":memory:")
    try:
        conn.enable_load_extension(True)
        sqlite_vec.load(conn)
        conn.enable_load_extension(False)
        conn.execute("CREATE VIRTUAL TABLE vd_doctor_smoke_test USING vec0(embedding float[4])")
        vec = sqlite_vec.serialize_float32([0.1, 0.2, 0.3, 0.4])
        conn.execute(
            "INSERT INTO vd_doctor_smoke_test(rowid, embedding) VALUES (1, ?)", (vec,)
        )
        row = conn.execute(
            "SELECT rowid, distance FROM vd_doctor_smoke_test "
            "WHERE embedding MATCH ? ORDER BY distance LIMIT 1",
            (vec,),
        ).fetchone()
        if row is None or row[0] != 1:
            raise RuntimeError(f"unexpected smoke-test round-trip result: {row!r}")
    except Exception as exc:
        return (
            CheckResult(
                "sqlite_vec_smoke_test",
                CheckStatus.WARN,
                f"sqlite-vec failed its Apple Silicon smoke test ({exc}); "
                "falling back to the NumPy vector-search backend.",
                required=False,
            ),
            "numpy",
        )
    finally:
        conn.close()

    return (
        CheckResult(
            "sqlite_vec_smoke_test",
            CheckStatus.PASS,
            "sqlite-vec loaded and passed a real insert+KNN round-trip.",
        ),
        "sqlite_vec",
    )


def check_pymupdf() -> CheckResult:
    try:
        import pymupdf as fitz  # 'import fitz' is deprecated and prints to
        # stdout on newer pymupdf releases — see pipeline/extract/pdf.py.

        doc = fitz.open()
        doc.new_page()
        doc[0].get_text()
        doc.close()
    except Exception as exc:
        return CheckResult("pymupdf", CheckStatus.FAIL, f"PyMuPDF smoke test failed: {exc}")
    return CheckResult(
        "pymupdf", CheckStatus.PASS, "PyMuPDF import + open/new-page/get_text round-trip OK."
    )


def check_python_docx() -> CheckResult:
    try:
        from docx import Document

        doc = Document()
        doc.add_paragraph("smoke test")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reopened = Document(buf)
        if reopened.paragraphs[0].text != "smoke test":
            raise RuntimeError("round-trip text mismatch")
    except Exception as exc:
        return CheckResult("python_docx", CheckStatus.FAIL, f"python-docx smoke test failed: {exc}")
    return CheckResult(
        "python_docx", CheckStatus.PASS, "python-docx create/save/reopen round-trip OK."
    )


def check_rapidocr() -> CheckResult:
    """Real (not stubbed) smoke test: run RapidOCR against a tiny
    self-generated image. Non-required — OCR is only needed once a user
    actually indexes scanned/image content, and the 'ocr' extra is optional.
    """
    try:
        import io

        from PIL import Image, ImageDraw

        from vectordrive.pipeline.ocr.base import OcrConfig
        from vectordrive.pipeline.ocr.rapidocr_engine import RapidOcrEngine

        img = Image.new("RGB", (200, 60), color="white")
        ImageDraw.Draw(img).text((5, 20), "doctor check", fill="black")
        buf = io.BytesIO()
        img.save(buf, format="PNG")

        result = RapidOcrEngine().run(buf.getvalue(), OcrConfig(lang="en"))
        if not result.success:
            return CheckResult(
                "rapidocr_smoke_test",
                CheckStatus.WARN,
                f"RapidOCR ran but reported failure: {result.error_message}",
                required=False,
            )
    except Exception as exc:
        return CheckResult(
            "rapidocr_smoke_test",
            CheckStatus.WARN,
            f"RapidOCR smoke test failed: {exc}. Install the 'ocr' extra: "
            "pip install -e '.[ocr]'.",
            required=False,
        )
    return CheckResult(
        "rapidocr_smoke_test",
        CheckStatus.PASS,
        "RapidOCR (PP-OCRv5 mobile, onnxruntime) round-trip OK.",
        required=False,
    )


def _import_paddle_and_paddleocr_vl():
    """Indirection point so tests can simulate an unavailable PaddleOCR-VL
    without needing it actually uninstalled.
    """
    import paddle
    from paddleocr import PaddleOCRVL  # noqa: F401 — import-only check

    return paddle


def check_paddleocr_vl_capability() -> CheckResult:
    """Lightweight import + version check only — NOT a full model load.

    A full round-trip (real PaddleOCR-VL-1.6 inference via the 'native'
    PaddlePaddle-CPU backend) was separately verified by hand during M5
    development; loading the 0.9B-param model takes ~50s cold, which would
    make every `doctor` run painfully slow if done here. This check just
    confirms the dependency chain (paddlepaddle>=3.2.1, paddleocr) is
    importable.
    """
    try:
        paddle = _import_paddle_and_paddleocr_vl()
    except Exception as exc:
        return CheckResult(
            "paddleocr_vl_capability",
            CheckStatus.WARN,
            f"PaddleOCR-VL not available ({exc}). Optional tier 3, for difficult "
            "scans/tables/handwriting/low-confidence pages. Install: "
            "pip install -e '.[paddleocr-vl]'.",
            required=False,
        )
    version = getattr(paddle, "__version__", "unknown")
    return CheckResult(
        "paddleocr_vl_capability",
        CheckStatus.PASS,
        f"paddlepaddle {version} and paddleocr importable (native CPU backend). "
        "Import/version check only — a real end-to-end round-trip was verified "
        "by hand during development.",
        required=False,
    )


def check_tesseract() -> CheckResult:
    """Informational only — Tesseract is optional and never the primary
    engine (requirement 5); its absence is not a problem.
    """
    import shutil

    path = shutil.which("tesseract")
    if path is None:
        return CheckResult(
            "tesseract_optional",
            CheckStatus.WARN,
            "tesseract binary not found (optional, last-resort fallback only — "
            "never the primary OCR engine). Install with `brew install tesseract` "
            "if you want it available.",
            required=False,
        )
    return CheckResult(
        "tesseract_optional",
        CheckStatus.PASS,
        f"tesseract binary found at {path}.",
        required=False,
    )


def check_ollama(endpoint: str, model: str) -> CheckResult:
    """Reachability + model-presence check only (no embedding call — the
    embedding stage itself is a later milestone). Non-required: v0.1 doesn't
    embed anything yet, so an unreachable Ollama shouldn't fail `doctor`.
    """
    try:
        resp = httpx.get(f"{endpoint}/api/tags", timeout=3.0)
        resp.raise_for_status()
    except Exception as exc:
        return CheckResult(
            "ollama_reachable",
            CheckStatus.WARN,
            f"Could not reach Ollama at {endpoint} ({exc}). Start Ollama "
            "before the embedding/search milestones.",
            required=False,
        )

    names = [m.get("name", "") for m in resp.json().get("models", [])]
    if model not in names:
        return CheckResult(
            "ollama_reachable",
            CheckStatus.WARN,
            f"Ollama reachable at {endpoint} but model '{model}' not found "
            f"(have: {', '.join(names) or 'none'}). Run `ollama pull {model}`.",
            required=False,
        )
    return CheckResult(
        "ollama_reachable",
        CheckStatus.PASS,
        f"Ollama reachable at {endpoint}; model '{model}' present.",
        required=False,
    )


def find_repo_root(start: Path | None = None) -> Path | None:
    """Walk upward from `start` (default: this file) for a pyproject.toml.

    Returns None when there is no source repository to find: a frozen
    (PyInstaller) app bundle carries no `pyproject.toml`, and the process
    working directory in a packaged launch is often ``/`` — which must
    never be mistaken for the repo root (every path is "inside" ``/``).
    The old ``Path.cwd()`` fallback did exactly that. A None result tells
    `check_data_dir_outside_repo` the check is not applicable.
    """
    if getattr(sys, "frozen", False):
        return None
    current = (start or Path(__file__)).resolve()
    for candidate in [current, *current.parents]:
        if (candidate / "pyproject.toml").exists():
            return candidate
    return None


def check_data_dir_outside_repo(repo_root: Path | None) -> CheckResult:
    from vectordrive.config import get_data_home

    data_home = get_data_home()
    if repo_root is None:
        # Packaged app (or any run with no source checkout): there is no
        # repository the data dir could pollute, so this development-only
        # guard passes trivially rather than inventing a bogus root.
        return CheckResult(
            "data_dir_outside_repo",
            CheckStatus.PASS,
            f"Data directory {data_home} — no source repository to check against "
            "(packaged application).",
        )
    if path_is_inside(data_home, repo_root):
        return CheckResult(
            "data_dir_outside_repo",
            CheckStatus.FAIL,
            f"VECTORDRIVE_HOME resolved to {data_home}, which is inside the "
            f"repo ({repo_root}). Set VECTORDRIVE_HOME to a path outside the repo.",
        )
    return CheckResult(
        "data_dir_outside_repo",
        CheckStatus.PASS,
        f"Data directory {data_home} is outside the repo.",
    )


def run_doctor(config: Config | None = None) -> list[CheckResult]:
    cfg = config or load_config()
    results: list[CheckResult] = []

    results.append(check_python_and_sqlite())
    vec_result, vector_backend = check_sqlite_vec_smoke_test()
    results.append(vec_result)
    results.append(check_ollama(cfg.ollama_endpoint, cfg.embedding_model))
    results.append(check_pymupdf())
    results.append(check_python_docx())
    results.append(check_data_dir_outside_repo(find_repo_root()))
    results.append(check_rapidocr())
    results.append(check_paddleocr_vl_capability())
    results.append(check_tesseract())

    cfg.vector_backend = vector_backend
    save_config(cfg)

    return results


def doctor_exit_code(results: list[CheckResult]) -> int:
    return 1 if any(r.status is CheckStatus.FAIL and r.required for r in results) else 0
